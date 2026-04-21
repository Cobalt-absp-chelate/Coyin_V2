from __future__ import annotations

import shutil
import tempfile
import zipfile
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Iterable
from xml.etree import ElementTree

import docx
from markdown_it import MarkdownIt
from pypdf import PdfReader

from coyin.core.common import hash_file, now_iso, short_id, slugify
from coyin.core.documents.models import (
    DocumentBlock,
    DocumentDescriptor,
    DocumentKind,
    DocumentSnapshot,
    OutlineItem,
)

try:
    import win32com.client  # type: ignore[import-untyped]
except Exception:  # pragma: no cover - optional dependency
    win32com = None


def _excerpt(text: str) -> str:
    compact = " ".join(text.split())
    return compact[:220]


class DocumentAdapter(ABC):
    kind: DocumentKind
    suffixes: tuple[str, ...]

    @abstractmethod
    def create_descriptor(self, path: Path) -> DocumentDescriptor:
        raise NotImplementedError

    @abstractmethod
    def load_snapshot(self, descriptor: DocumentDescriptor) -> DocumentSnapshot:
        raise NotImplementedError

    def load_reader_snapshot(self, descriptor: DocumentDescriptor) -> DocumentSnapshot:
        return self.load_snapshot(descriptor)

    def _base_descriptor(self, path: Path, title: str) -> DocumentDescriptor:
        return DocumentDescriptor(
            document_id=short_id("doc"),
            title=title,
            path=str(path),
            kind=self.kind.value,
            added_at=now_iso(),
            fingerprint=hash_file(path),
        )


class PdfAdapter(DocumentAdapter):
    kind = DocumentKind.PDF
    suffixes = (".pdf",)

    def create_descriptor(self, path: Path) -> DocumentDescriptor:
        reader = PdfReader(str(path))
        metadata = reader.metadata or {}
        title = str(metadata.get("/Title") or path.stem).strip() or path.stem
        authors = [part.strip() for part in str(metadata.get("/Author") or "").split(",") if part.strip()]
        year = ""
        descriptor = self._base_descriptor(path, title)
        descriptor.authors = authors
        descriptor.year = year
        descriptor.excerpt = _excerpt(self._extract_text(reader))
        descriptor.metadata = {"page_count": len(reader.pages)}
        return descriptor

    def _extract_text(self, reader: PdfReader) -> str:
        fragments: list[str] = []
        for page in reader.pages[:20]:
            try:
                fragments.append(page.extract_text() or "")
            except Exception:
                continue
        return "\n".join(fragments)

    def load_snapshot(self, descriptor: DocumentDescriptor) -> DocumentSnapshot:
        reader = PdfReader(descriptor.path)
        outline = _pdf_outline(reader)
        blocks: list[DocumentBlock] = []
        raw_parts: list[str] = []
        for index, page in enumerate(reader.pages):
            try:
                text = page.extract_text() or ""
            except Exception:
                text = ""
            raw_parts.append(text)
            blocks.append(
                DocumentBlock(
                    block_id=f"page-{index + 1}",
                    kind="page",
                    text=text,
                    page=index,
                    meta={"label": f"第 {index + 1} 页"},
                )
            )
        return DocumentSnapshot(
            document_id=descriptor.document_id,
            raw_text="\n".join(raw_parts),
            blocks=blocks,
            outline=outline,
            page_count=len(reader.pages),
            meta={"source": "pdf"},
        )

    def load_reader_snapshot(self, descriptor: DocumentDescriptor) -> DocumentSnapshot:
        reader = PdfReader(descriptor.path)
        return DocumentSnapshot(
            document_id=descriptor.document_id,
            raw_text="",
            blocks=[],
            outline=_pdf_outline(reader),
            page_count=len(reader.pages),
            meta={"source": "pdf", "staged": True},
        )


def _pdf_outline(reader: PdfReader) -> list[OutlineItem]:
    try:
        outline = reader.outline
    except Exception:
        return []

    def convert(items: Iterable) -> list[OutlineItem]:
        result: list[OutlineItem] = []
        for item in items:
            if isinstance(item, list):
                if result:
                    result[-1].children.extend(convert(item))
                continue
            title = getattr(item, "title", None) or str(item)
            page = None
            try:
                page = reader.get_destination_page_number(item)
            except Exception:
                page = None
            result.append(OutlineItem(title=title, page=page))
        return result

    return convert(outline)


class PlainTextAdapter(DocumentAdapter):
    kind = DocumentKind.TEXT
    suffixes = (".txt",)

    def create_descriptor(self, path: Path) -> DocumentDescriptor:
        text = path.read_text(encoding="utf-8", errors="ignore")
        title = next((line.strip() for line in text.splitlines() if line.strip()), path.stem)
        descriptor = self._base_descriptor(path, title)
        descriptor.excerpt = _excerpt(text)
        return descriptor

    def load_snapshot(self, descriptor: DocumentDescriptor) -> DocumentSnapshot:
        text = Path(descriptor.path).read_text(encoding="utf-8", errors="ignore")
        blocks = [
            DocumentBlock(block_id=f"line-{index}", kind="paragraph", text=line)
            for index, line in enumerate(text.splitlines(), start=1)
            if line.strip()
        ]
        return DocumentSnapshot(document_id=descriptor.document_id, raw_text=text, blocks=blocks)


class MarkdownAdapter(DocumentAdapter):
    kind = DocumentKind.MARKDOWN
    suffixes = (".md", ".markdown")

    def create_descriptor(self, path: Path) -> DocumentDescriptor:
        text = path.read_text(encoding="utf-8", errors="ignore")
        title = path.stem
        for line in text.splitlines():
            if line.strip().startswith("#"):
                title = line.lstrip("#").strip()
                break
        descriptor = self._base_descriptor(path, title)
        descriptor.excerpt = _excerpt(text)
        return descriptor

    def load_snapshot(self, descriptor: DocumentDescriptor) -> DocumentSnapshot:
        text = Path(descriptor.path).read_text(encoding="utf-8", errors="ignore")
        blocks: list[DocumentBlock] = []
        outline: list[OutlineItem] = []
        for index, line in enumerate(text.splitlines(), start=1):
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("#"):
                depth = len(stripped) - len(stripped.lstrip("#"))
                title = stripped[depth:].strip()
                outline.append(OutlineItem(title=title, anchor=f"line-{index}"))
                blocks.append(DocumentBlock(block_id=f"line-{index}", kind=f"heading-{depth}", text=title))
            else:
                blocks.append(DocumentBlock(block_id=f"line-{index}", kind="paragraph", text=stripped))
        return DocumentSnapshot(document_id=descriptor.document_id, raw_text=text, blocks=blocks, outline=outline)

    def to_html(self, text: str) -> str:
        return MarkdownIt().render(text)


class DocxAdapter(DocumentAdapter):
    kind = DocumentKind.DOCX
    suffixes = (".docx",)

    def create_descriptor(self, path: Path) -> DocumentDescriptor:
        package = docx.Document(str(path))
        core = package.core_properties
        title = core.title or path.stem
        authors = [core.author] if core.author else []
        text = "\n".join(paragraph.text for paragraph in package.paragraphs)
        descriptor = self._base_descriptor(path, title)
        descriptor.authors = [author for author in authors if author]
        descriptor.excerpt = _excerpt(text)
        descriptor.metadata = {
            "subject": core.subject or "",
            "keywords": core.keywords or "",
        }
        return descriptor

    def load_snapshot(self, descriptor: DocumentDescriptor) -> DocumentSnapshot:
        package = docx.Document(descriptor.path)
        blocks: list[DocumentBlock] = []
        raw_parts: list[str] = []
        for index, paragraph in enumerate(package.paragraphs, start=1):
            text = paragraph.text.strip()
            if not text:
                continue
            raw_parts.append(text)
            style = paragraph.style.name.lower() if paragraph.style else "paragraph"
            blocks.append(DocumentBlock(block_id=f"p-{index}", kind=style, text=text))
        return DocumentSnapshot(document_id=descriptor.document_id, raw_text="\n".join(raw_parts), blocks=blocks)


class DocBinaryAdapter(DocumentAdapter):
    kind = DocumentKind.DOC
    suffixes = (".doc",)

    def create_descriptor(self, path: Path) -> DocumentDescriptor:
        descriptor = self._base_descriptor(path, path.stem)
        descriptor.excerpt = "已导入旧版 Word 文档。若本机安装了 Word，将在打开时自动转换提取。"
        descriptor.metadata = {"legacy_word": True}
        return descriptor

    def load_snapshot(self, descriptor: DocumentDescriptor) -> DocumentSnapshot:
        converted = self._convert_to_docx(Path(descriptor.path))
        if converted:
            temp_descriptor = DocumentDescriptor(
                document_id=descriptor.document_id,
                title=descriptor.title,
                path=str(converted),
                kind=DocumentKind.DOCX.value,
            )
            return DocxAdapter().load_snapshot(temp_descriptor)
        return DocumentSnapshot(
            document_id=descriptor.document_id,
            raw_text="",
            blocks=[DocumentBlock(block_id="legacy", kind="notice", text="当前环境无法直接解析该 .doc 文件。")],
        )

    def _convert_to_docx(self, source: Path) -> Path | None:
        if win32com is None:
            return None
        temp_dir = Path(tempfile.mkdtemp(prefix="coyin_doc_"))
        target = temp_dir / f"{slugify(source.stem)}.docx"
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            document = word.Documents.Open(str(source))
            document.SaveAs(str(target), FileFormat=16)
            document.Close(False)
            word.Quit()
            return target
        except Exception:
            return None


class LatexAdapter(DocumentAdapter):
    kind = DocumentKind.LATEX
    suffixes = (".tex",)

    def create_descriptor(self, path: Path) -> DocumentDescriptor:
        text = path.read_text(encoding="utf-8", errors="ignore")
        title = next((line for line in text.splitlines() if "\\title" in line), path.stem)
        descriptor = self._base_descriptor(path, path.stem if title == path.stem else title)
        descriptor.excerpt = _excerpt(text)
        return descriptor

    def load_snapshot(self, descriptor: DocumentDescriptor) -> DocumentSnapshot:
        text = Path(descriptor.path).read_text(encoding="utf-8", errors="ignore")
        blocks = [
            DocumentBlock(block_id=f"line-{index}", kind="latex", text=line)
            for index, line in enumerate(text.splitlines(), start=1)
        ]
        return DocumentSnapshot(document_id=descriptor.document_id, raw_text=text, blocks=blocks)


class DraftAdapter(DocumentAdapter):
    kind = DocumentKind.DRAFT
    suffixes = (".cydraft",)

    def create_descriptor(self, path: Path) -> DocumentDescriptor:
        text = path.read_text(encoding="utf-8", errors="ignore")
        descriptor = self._base_descriptor(path, path.stem)
        descriptor.excerpt = _excerpt(text)
        return descriptor

    def load_snapshot(self, descriptor: DocumentDescriptor) -> DocumentSnapshot:
        text = Path(descriptor.path).read_text(encoding="utf-8", errors="ignore")
        return DocumentSnapshot(
            document_id=descriptor.document_id,
            raw_text=text,
            blocks=[DocumentBlock(block_id="draft", kind="richtext", text=text)],
        )


class DocumentAdapterRegistry:
    def __init__(self):
        self.adapters: list[DocumentAdapter] = [
            PdfAdapter(),
            MarkdownAdapter(),
            PlainTextAdapter(),
            DocxAdapter(),
            DocBinaryAdapter(),
            LatexAdapter(),
            DraftAdapter(),
        ]

    def for_path(self, path: Path) -> DocumentAdapter | None:
        suffix = path.suffix.lower()
        for adapter in self.adapters:
            if suffix in adapter.suffixes:
                return adapter
        return None
