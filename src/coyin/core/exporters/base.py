from __future__ import annotations

import re
from pathlib import Path

from PySide6.QtGui import QTextDocument
from PySide6.QtPrintSupport import QPrinter
from docx import Document as DocxDocument


class DraftExporter:
    def export_pdf(self, html: str, target: Path) -> None:
        document = QTextDocument()
        document.setHtml(html)
        printer = QPrinter(QPrinter.PrinterMode.HighResolution)
        printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
        printer.setOutputFileName(str(target))
        document.print_(printer)

    def export_docx(self, title: str, html: str, target: Path) -> None:
        document = DocxDocument()
        document.add_heading(title, level=1)
        text = self.html_to_plain_blocks(html)
        for block in text:
            if block.startswith("# "):
                document.add_heading(block[2:], level=2)
            else:
                document.add_paragraph(block)
        document.save(str(target))

    def export_markdown(self, html: str, target: Path) -> None:
        blocks = self.html_to_plain_blocks(html)
        target.write_text("\n\n".join(blocks), encoding="utf-8")

    def html_to_plain_blocks(self, html: str) -> list[str]:
        cleaned = re.sub(r"<br\s*/?>", "\n", html, flags=re.I)
        cleaned = re.sub(r"</p>", "\n\n", cleaned, flags=re.I)
        cleaned = re.sub(r"<[^>]+>", "", cleaned)
        cleaned = cleaned.replace("&nbsp;", " ").replace("&amp;", "&")
        return [block.strip() for block in cleaned.split("\n\n") if block.strip()]
